import os
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_excel_report(results, output_dir="tests/reports"):
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filepath = os.path.join(output_dir, f"test-report_{timestamp}.xlsx")
    
    df = pd.DataFrame(results)
    
    # Ensure columns exist even if results are empty
    expected_cols = ["Test ID", "Module", "Flow Name", "Test Description", 
                     "Request Payload", "Expected Result", "Actual Result", 
                     "Status", "Error Message", "Execution Time", "Timestamp"]
    for col in expected_cols:
        if col not in df.columns:
            df[col] = ""
            
    df = df[expected_cols]
    
    # Write to Excel with formatting
    with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Test Results')
        workbook = writer.book
        worksheet = writer.sheets['Test Results']
        
        # Adjust column widths
        for i, col in enumerate(df.columns):
            max_len = max(df[col].astype(str).map(len).max(), len(col)) + 2
            worksheet.column_dimensions[chr(65 + i)].width = min(max_len, 50)
            
    print(f"Excel report generated: {filepath}")
    return filepath

def generate_docx_report(results, output_dir="tests/reports"):
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filepath = os.path.join(output_dir, f"integration-summary_{timestamp}.docx")
    
    doc = Document()
    
    # Title
    title = doc.add_heading('PortSync End-to-End Test Execution Summary', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Execution Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Statistics
    total = len(results)
    passed = sum(1 for r in results if r.get('Status') == 'PASS')
    failed = total - passed
    
    doc.add_heading('Execution Statistics', level=1)
    stats_table = doc.add_table(rows=2, cols=3)
    stats_table.style = 'Table Grid'
    hdr_cells = stats_table.rows[0].cells
    hdr_cells[0].text = 'Total Tests'
    hdr_cells[1].text = 'Passed'
    hdr_cells[2].text = 'Failed'
    
    row_cells = stats_table.rows[1].cells
    row_cells[0].text = str(total)
    row_cells[1].text = str(passed)
    row_cells[2].text = str(failed)
    
    # Failures section
    doc.add_heading('Failed Test Details', level=1)
    if failed == 0:
        doc.add_paragraph("All tests passed successfully! No failures to report.")
    else:
        for r in results:
            if r.get('Status') != 'PASS':
                doc.add_heading(r.get('Test ID', 'Unknown Test'), level=2)
                doc.add_paragraph(f"Module: {r.get('Module', '')}")
                doc.add_paragraph(f"Description: {r.get('Test Description', '')}")
                doc.add_paragraph(f"Error: {r.get('Error Message', '')}")
                if 'screenshot' in r and r['screenshot']:
                    try:
                        doc.add_picture(r['screenshot'], width=Inches(6.0))
                    except Exception as e:
                        doc.add_paragraph(f"(Screenshot missing: {str(e)})")
                doc.add_paragraph("-" * 40)
    
    doc.save(filepath)
    print(f"Word document report generated: {filepath}")
    return filepath
